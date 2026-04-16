from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.section import Section
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class DocxRedactionHandler:
    def __init__(self, text_redactor) -> None:
        self.text_redactor = text_redactor

    def process(self, input_path: Path, output_path: Path, options) -> Path:
        document = Document(str(input_path))

        for paragraph in self._iter_all_paragraphs(document):
            self._replace_paragraph_text(paragraph, options)

        document.save(str(output_path))
        return output_path

    def _iter_all_paragraphs(self, document: DocumentObject) -> Iterable[Paragraph]:
        yield from document.paragraphs
        for table in document.tables:
            yield from self._iter_table_paragraphs(table)
        for section in document.sections:
            yield from self._iter_section_paragraphs(section)

    def _iter_section_paragraphs(self, section: Section) -> Iterable[Paragraph]:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from self._iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table: Table) -> Iterable[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                yield from self._iter_cell_paragraphs(cell)

    def _iter_cell_paragraphs(self, cell: _Cell) -> Iterable[Paragraph]:
        for paragraph in cell.paragraphs:
            yield paragraph
        for table in cell.tables:
            yield from self._iter_table_paragraphs(table)

    def _replace_paragraph_text(self, paragraph: Paragraph, options) -> None:
        if not paragraph.runs:
            return

        original_text = "".join(run.text for run in paragraph.runs)
        if not original_text.strip():
            return

        redacted_text = self.text_redactor.redact_text(original_text, options)
        if redacted_text == original_text:
            return

        run_lengths = [len(run.text) for run in paragraph.runs]
        cursor = 0
        for index, run in enumerate(paragraph.runs):
            if index < len(paragraph.runs) - 1:
                next_cursor = cursor + run_lengths[index]
                run.text = redacted_text[cursor:next_cursor]
                cursor = next_cursor
            else:
                run.text = redacted_text[cursor:]
